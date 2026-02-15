# components/ai_suggester.py - env-driven GPT, cache-safe doc builders, VP critique loop

import os, json, hashlib, tempfile
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from slugify import slugify

load_dotenv()

def _get_openai_key():
    """Get OpenAI API key from Streamlit secrets or environment."""
    # Try Streamlit secrets first (for Streamlit Cloud)
    try:
        import streamlit as st
        if hasattr(st, 'secrets') and 'OPENAI_API_KEY' in st.secrets:
            return st.secrets['OPENAI_API_KEY']
    except Exception:
        pass
    # Fall back to environment variable
    return os.getenv("OPENAI_API_KEY")

_api_key = _get_openai_key()
client = OpenAI(api_key=_api_key) if _api_key else None

MODEL_MAIN = os.getenv("OPENAI_MODEL_MAIN", "gpt-4o")
MODEL_MINI = os.getenv("OPENAI_MODEL_SCREENER", "gpt-4o-mini")
CACHE_PATH = "gpt_suggestion_cache.json"


def _load_cache():
    if os.path.exists(CACHE_PATH):
        try:
            with open(CACHE_PATH, "r", encoding="utf-8") as f:
                content = f.read().strip()
                return json.loads(content) if content else {}
        except Exception:
            return {}
    return {}


_sugg_cache = _load_cache()


def _save_cache():
    with open(CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(_sugg_cache, f, indent=2)


def cache_and_return(key, value):
    _sugg_cache[key] = value
    _save_cache()
    return value


def clean_gpt_input(text, max_words=1000):
    return " ".join((text or "").strip().split()[:max_words])


def should_fallback_to_signal_brief(text):
    t = (text or "").strip()
    return len(t) < 50 or len(t.split()) < 10


def safe_file_path(base_name, prefix="insight"):
    filename = slugify(f"{prefix}-{base_name}")[:64] + ".docx"
    return os.path.join(tempfile.gettempdir(), filename)


def write_docx(content, heading):
    doc = Document()
    doc.add_heading(heading, level=1)
    for line in (content or "").split("\n"):
        if line.strip().endswith(":"):
            doc.add_heading(line.strip(), level=2)
        else:
            doc.add_paragraph(line.strip())
    return doc


def build_metadata_block(brand, trend_context=None, competitor_context=None, meta_fields=None, insight=None):
    ctx = [
        "═══════════════════════════════════════",
        "CONTEXTUAL METADATA (use to inform all sections)",
        "═══════════════════════════════════════",
        f"Brand/Platform: {brand}",
        "Primary Objective: Improve trust, conversion, reduce friction, or increase GMV.",
    ]
    if insight:
        ctx.append(f"Sentiment: {insight.get('brand_sentiment', 'Unknown')}")
        ctx.append(f"Persona: {insight.get('persona', 'General')}")
        ctx.append(f"Journey Stage: {insight.get('journey_stage', 'Unknown')}")
        ctx.append(f"Topic Focus: {', '.join(insight.get('topic_focus_list', insight.get('topic_focus', [])) or ['General'])}")
        ctx.append(f"Effort Estimate: {insight.get('effort', 'Unknown')}")
        ctx.append(f"Severity Score: {insight.get('severity_score', 'N/A')}")
        ctx.append(f"PM Priority Score: {insight.get('pm_priority_score', 'N/A')}")
        if insight.get('_payment_issue'):
            ctx.append(f"⚠️ Payment Friction Flag: Yes — Types: {', '.join(insight.get('payment_issue_types', []))}")
        if insight.get('_upi_flag'):
            ctx.append("⚠️ UPI/Unpaid Item Flag: Yes (seller impact)")
        if insight.get('_high_end_flag'):
            ctx.append("💎 High-ASP Flag: Yes (high-value transaction context)")
        if insight.get('opportunity_tag'):
            ctx.append(f"Opportunity Type: {insight.get('opportunity_tag')}")
        if insight.get('mentions_competitor'):
            ctx.append(f"Competitors Mentioned: {', '.join(insight.get('mentions_competitor', []))}")
        if insight.get('mentions_ecosystem_partner'):
            ctx.append(f"Partners Mentioned: {', '.join(insight.get('mentions_ecosystem_partner', []))}")
    if trend_context:
        ctx.append(f"Trend Signal: {trend_context}")
    if competitor_context:
        ctx.append(f"Competitor Context: {competitor_context}")
    if meta_fields:
        for k, v in meta_fields.items():
            ctx.append(f"{k}: {v}")
    ctx.append("═══════════════════════════════════════")
    ctx.append("IMPORTANT: Fill in ALL sections with specific, actionable content. Do NOT leave placeholders like [TBD] or [insert here]. Make reasonable assumptions based on the signal.")
    return "\n".join(ctx)


def generate_exec_summary():
    return (
        "\n\n---\n\n**Executive TL;DR**\n"
        "- What: [summary]\n- Why it matters: [impact]\n- What decision is needed: [action]"
    )


# ------------------------------
# Core chat helper (safe when no API key)
# ------------------------------
def _chat(model, system, user, max_completion_tokens=2000, temperature=0.3):
    """
    Wrapper around chat.completions.create that uses max_completion_tokens
    so it works with newer models like gpt-5.1 and newer models.

    """
    if client is None:
        # Offline or no key: return a stub so the pipeline does not break
        return f"[LLM disabled] {system}\n\n{user[:800]}"
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=temperature,
        max_completion_tokens=max_completion_tokens,
    )
    return (resp.choices[0].message.content or "").strip()


def generate_gpt_doc(prompt, title, max_tokens=4000):
    try:
        draft = _chat(
            MODEL_MAIN,
            title,
            clean_gpt_input(prompt, max_words=2000),
            max_completion_tokens=max_tokens,
            temperature=0.3,
        )
        # If running in fallback or offline, skip critique loop
        if draft.startswith("[LLM disabled]"):
            return draft

        critique = (
            "You are a VP of Product reviewing this document. Your job:\n"
            "1. Ensure ALL sections are filled with specific, concrete content (no placeholders)\n"
            "2. Add quantitative estimates where missing (TAM, effort in weeks, % impact)\n"
            "3. Tighten structure, remove fluff, improve specificity\n"
            "4. Ensure success metrics are measurable and time-bound\n"
            "5. Add any missing critical sections\n\n"
            "Rewrite the full document with improvements:\n\n"
            f"{draft}"
        )
        return _chat(
            MODEL_MAIN,
            "You are a critical VP of Product ensuring document completeness.",
            clean_gpt_input(critique, max_words=3000),
            max_completion_tokens=max_tokens,
            temperature=0.3,
        )
    except Exception as e:
        return f"⚠️ GPT Error: {e}"


def generate_pm_ideas(text, brand="eBay"):
    key = hashlib.md5(f"{text}_{brand}".encode()).hexdigest()
    if key in _sugg_cache:
        return _sugg_cache[key]

    prompt = (
        "You are a senior PM at a marketplace like eBay.\n"
        "Generate 3 concise, concrete product suggestions to improve trust or conversion.\n\n"
        f"Feedback:\n{text}\n\nBrand: {brand}"
    )
    try:
        ideas = _chat(
            MODEL_MINI,
            "Generate actionable, concrete product improvement ideas.",
            prompt,
            max_completion_tokens=320,
            temperature=0.2,
        )
        # If offline stub, just return one idea with trimmed prompt context
        if ideas.startswith("[LLM disabled]"):
            return cache_and_return(key, [ideas[:240]])

        lines = [l.strip("-• ").strip() for l in ideas.split("\n") if l.strip()]
        return cache_and_return(key, (lines[:3] or [ideas]))
    except Exception as e:
        return [f"[GPT error: {e}]"]


# --- compat: simple batch wrapper so precompute_insights.py can import it ---
def generate_pm_ideas_batch(texts, brand="eBay"):
    """
    Batch-safe wrapper. Returns a list of idea lists (one list per input text).
    Falls back gracefully if any single suggestion fails.
    """
    results = []
    for t in (texts or []):
        try:
            results.append(generate_pm_ideas(text=t, brand=brand))
        except Exception as e:
            results.append([f"[GPT error: {e}]"])
    return results


def _maybe_brief(text, brand, base_filename):
    prompt = (
        "Turn this brief signal into a 1-page internal summary for product leadership.\n\n"
        f"{text}\n\nBrand: {brand}\n\nSections:\n"
        "- Observation\n- Hypothesis\n- Strategic Importance\n- Potential Impact\n- Suggested Next Steps\n- Open Questions"
    )
    content = generate_gpt_doc(prompt, "You summarize vague signals into a strategic brief.")
    doc = write_docx(content, "Strategic Signal Brief")
    path = safe_file_path(base_filename, prefix="brief")
    doc.save(path)
    return path


def generate_prd_docx(text, brand, base_filename, trend_context=None, competitor_context=None, meta_fields=None, insight=None):
    if should_fallback_to_signal_brief(text):
        return _maybe_brief(text, brand, base_filename)
    meta = build_metadata_block(brand, trend_context, competitor_context, meta_fields, insight)
    prompt = f"""You are a senior product manager at {brand}. Write a comprehensive, GTM-ready Product Requirements Document (PRD) based on the following user signal/feedback.

USER SIGNAL:
{text}

{meta}

DOCUMENT STRUCTURE (fill in ALL sections with specific, actionable content):

1. EXECUTIVE SUMMARY (3 bullets max):
   - What: One sentence describing the proposed solution
   - Why: Business impact and user pain being addressed
   - Ask: What decision or resources are needed

2. PROBLEM STATEMENT:
   - Describe the user pain point in detail
   - Include verbatim quotes from the signal
   - Quantify the impact (estimate affected users, frequency, severity)

3. STRATEGIC CONTEXT:
   - How does this align with {brand}'s strategy?
   - Market trends supporting this investment
   - Competitive landscape (what are competitors doing?)

4. USER PERSONAS:
   - Primary persona (name, description, pain points, goals)
   - Secondary personas affected
   - Jobs to be done (JTBD) for each persona

5. CURRENT STATE VS FUTURE STATE:
   - Current user journey with pain points highlighted
   - Proposed future journey with improvements
   - Before/after comparison

6. PROPOSED SOLUTION:
   - High-level solution description
   - Key features/capabilities (bullet list with descriptions)
   - Out of scope (what we're NOT doing)

7. REQUIREMENTS:
   - Functional requirements (must-have)
   - Non-functional requirements (performance, security, accessibility)
   - Technical dependencies

8. SUCCESS METRICS:
   - Primary KPI with target (e.g., "Reduce support tickets by 25% within 90 days")
   - Secondary metrics
   - How will we measure success?

9. EFFORT & TIMELINE:
   - T-shirt size estimate (S/M/L/XL)
   - Estimated weeks to MVP
   - Key milestones

10. RISKS & MITIGATIONS:
    - Technical risks
    - Business risks
    - Dependencies on other teams

11. HYPOTHESIS & EXPERIMENT:
    - Hypothesis statement
    - Suggested A/B test or pilot approach
    - Success criteria for experiment

12. STAKEHOLDERS & APPROVALS:
    - Product owner
    - Engineering lead
    - Design lead
    - Other stakeholders

13. OPEN QUESTIONS:
    - List unresolved questions that need input

14. APPENDIX:
    - Related signals or supporting data
    - Competitive screenshots (if relevant)
    - Technical diagrams (placeholder)"""
    content = generate_gpt_doc(prompt, "You are a senior product manager writing a comprehensive PRD.", max_tokens=5000)
    doc = write_docx(content, "Product Requirements Document (PRD)")
    path = safe_file_path(base_filename, prefix="prd")
    doc.save(path)
    return path


def generate_brd_docx(text, brand, base_filename, trend_context=None, competitor_context=None, meta_fields=None, insight=None):
    if should_fallback_to_signal_brief(text):
        return _maybe_brief(text, brand, base_filename)
    meta = build_metadata_block(brand, trend_context, competitor_context, meta_fields, insight)
    prompt = f"""You are a senior business strategist at {brand}. Write a comprehensive Business Requirements Document (BRD) based on the following user signal/feedback.

USER SIGNAL:
{text}

{meta}

DOCUMENT STRUCTURE (fill in ALL sections with specific, actionable content):

1. EXECUTIVE SUMMARY:
   - What: One sentence describing the business opportunity
   - Why: Revenue/cost impact being addressed
   - Ask: Investment or decision needed

2. BUSINESS PROBLEM:
   - Current business pain point
   - Verbatim user quotes supporting the problem
   - Impact on key business metrics (GMV, conversion, retention, support costs)

3. MARKET OPPORTUNITY:
   - Total Addressable Market (TAM) - estimate with rationale
   - Serviceable Addressable Market (SAM)
   - Serviceable Obtainable Market (SOM)
   - Market trends supporting investment

4. COMPETITIVE ANALYSIS:
   - How competitors address this problem
   - Competitive advantage/disadvantage
   - Market positioning opportunity

5. STRATEGIC ALIGNMENT:
   - How this aligns with {brand}'s strategic priorities
   - Which company OKRs does this support?
   - Strategic bet assessment (risk × reward matrix)

6. AFFECTED USER SEGMENTS:
   - Primary segment (size, characteristics, value)
   - Secondary segments
   - User segment growth trends

7. BUSINESS SOLUTION:
   - High-level solution approach
   - Key capabilities required
   - Build vs. buy analysis

8. FINANCIAL ANALYSIS:
   - Revenue impact estimate (with assumptions)
   - Cost savings estimate
   - Implementation cost estimate
   - ROI calculation with payback period
   - NPV/IRR if applicable

9. LEGAL, COMPLIANCE & POLICY:
   - Regulatory considerations
   - Policy changes required
   - Privacy/data implications
   - Geographic considerations

10. IMPLEMENTATION APPROACH:
    - Phased rollout plan
    - Resource requirements
    - Timeline to value

11. SUCCESS METRICS & TARGETS:
    - Primary business KPI with target
    - Secondary metrics
    - Measurement methodology

12. RISKS & DEPENDENCIES:
    - Business risks
    - Technical dependencies
    - Market risks
    - Mitigation strategies

13. STAKEHOLDERS:
    - Executive sponsor
    - Business owner
    - Cross-functional partners
    - Approval chain

14. RECOMMENDATION:
    - Clear go/no-go recommendation
    - Investment ask
    - Expected return
    - Confidence level (High/Medium/Low)"""
    content = generate_gpt_doc(prompt, "You are a senior business strategist writing a comprehensive BRD.", max_tokens=5000)
    doc = write_docx(content, "Business Requirements Document (BRD)")
    path = safe_file_path(base_filename, prefix="brd")
    doc.save(path)
    return path


def generate_prfaq_docx(text, brand, base_filename, trend_context=None, competitor_context=None, meta_fields=None, insight=None):
    if should_fallback_to_signal_brief(text):
        return _maybe_brief(text, brand, base_filename)
    meta = build_metadata_block(brand, trend_context, competitor_context, meta_fields, insight)
    prompt = f"""You are a product marketing leader at {brand}. Write a comprehensive Amazon-style PRFAQ (Press Release + FAQ) based on the following user signal/feedback.

USER SIGNAL:
{text}

{meta}

DOCUMENT STRUCTURE (fill in ALL sections with specific, compelling content):

═══════════════════════════════════════
PART 1: PRESS RELEASE
═══════════════════════════════════════

HEADLINE:
- Attention-grabbing headline announcing the feature/product
- Should communicate the core benefit to users

SUBHEADLINE:
- One sentence expanding on the headline
- Include the target user and primary benefit

DATELINE & INTRO PARAGRAPH:
- City, Date — {brand} today announced...
- What is being launched and why it matters
- Who benefits from this

PROBLEM PARAGRAPH:
- Describe the customer problem being solved
- Use specific examples and pain points from the signal
- Make it relatable and urgent

SOLUTION PARAGRAPH:
- How the new feature/product solves the problem
- Key capabilities and benefits
- What makes this approach unique

CUSTOMER QUOTE (fictional but realistic):
- Create a realistic customer quote
- Include their name, role, and context
- Express genuine relief/satisfaction with the solution

EXECUTIVE QUOTE:
- Quote from a fictional VP/Director at {brand}
- Explain why this matters strategically
- Express commitment to the customer

HOW IT WORKS:
- Step-by-step explanation (3-5 steps)
- Clear, jargon-free language
- Highlight ease of use

AVAILABILITY:
- When/where this will be available
- Any phased rollout details
- How to access or sign up

CALL TO ACTION:
- What should users do next?
- Link placeholder for more information

═══════════════════════════════════════
PART 2: FREQUENTLY ASKED QUESTIONS
═══════════════════════════════════════

CUSTOMER FAQs (answer 6-8 questions):
- What is this feature and how does it work?
- Who is this for?
- How much does it cost?
- When will this be available?
- How is this different from what exists today?
- What if I have a problem?
- Will this work with [related feature]?
- How do I get started?

INTERNAL/STAKEHOLDER FAQs (answer 4-6 questions):
- Why are we building this now?
- What's the expected business impact?
- What are the key risks?
- What resources are required?
- How will we measure success?
- What's the competitive response risk?

OBJECTION HANDLING:
- List 3-5 likely objections from skeptics
- Provide compelling responses to each

═══════════════════════════════════════
PART 3: GTM READINESS CHECKLIST
═══════════════════════════════════════

- [ ] Customer research validated
- [ ] Competitive analysis complete
- [ ] Pricing/packaging defined
- [ ] Success metrics defined
- [ ] Support documentation ready
- [ ] Marketing assets prepared
- [ ] Sales enablement complete
- [ ] Legal/compliance approved
- [ ] Rollout plan finalized
- [ ] Rollback plan documented"""
    content = generate_gpt_doc(prompt, "You are a product marketing leader writing a comprehensive PRFAQ.", max_tokens=6000)
    doc = write_docx(content, "Product PRFAQ Document")
    path = safe_file_path(base_filename, prefix="faq")
    doc.save(path)
    return path


def generate_jira_bug_ticket(text, brand="eBay", insight=None):
    """Generate a detailed JIRA ticket based on insight type."""
    insight_type = insight.get("type_tag", "Bug") if insight else "Bug"
    subtag = insight.get("subtag", "General") if insight else "General"
    
    if insight_type == "Question":
        ticket_type = "Task"
        prompt = f"""Create a JIRA Task ticket for this user question/confusion:

USER FEEDBACK:
{text}

CONTEXT: {subtag} issue on {brand}

Generate a well-structured JIRA ticket with:

**Title:** [Clear, actionable title]

**Type:** Task

**Priority:** Medium

**Labels:** user-feedback, {subtag.lower().replace(' ', '-')}, documentation

**Description:**
## User Question
[Summarize what the user is confused about]

## Root Cause Analysis
[Why is this confusing? Is documentation unclear? Is the UX not intuitive?]

## Acceptance Criteria
- [ ] User can easily find answer to this question
- [ ] Documentation is updated (if applicable)
- [ ] UX improvements identified (if applicable)

## Suggested Actions
1. [First action]
2. [Second action]
3. [Third action]

## Related Areas
[List related features or documentation]"""

    elif insight_type == "Feature Request":
        ticket_type = "Story"
        prompt = f"""Create a JIRA Story ticket for this feature request:

USER FEEDBACK:
{text}

CONTEXT: {subtag} feature request for {brand}

Generate a well-structured JIRA ticket with:

**Title:** [User story format: As a [user], I want [feature] so that [benefit]]

**Type:** Story

**Priority:** Medium

**Labels:** feature-request, {subtag.lower().replace(' ', '-')}, user-feedback

**Description:**
## User Story
As a [persona], I want [feature description] so that [benefit].

## User Feedback (Verbatim)
> {text[:500]}

## Problem Statement
[What problem does this solve?]

## Proposed Solution
[High-level solution approach]

## Acceptance Criteria
- [ ] [Criteria 1]
- [ ] [Criteria 2]
- [ ] [Criteria 3]

## Business Value
[Why should we build this? What's the impact?]

## Technical Considerations
[Any known technical constraints or dependencies]"""

    else:  # Bug/Complaint
        ticket_type = "Bug"
        prompt = f"""Create a JIRA Bug ticket for this user complaint:

USER FEEDBACK:
{text}

CONTEXT: {subtag} issue on {brand}

Generate a well-structured JIRA ticket with:

**Title:** [{subtag}] [Clear bug description]

**Type:** Bug

**Priority:** High

**Labels:** bug, {subtag.lower().replace(' ', '-')}, user-reported

**Description:**
## Bug Summary
[One sentence describing the bug]

## User Report (Verbatim)
> {text[:500]}

## Steps to Reproduce
1. [Step 1]
2. [Step 2]
3. [Step 3]

## Expected Behavior
[What should happen]

## Actual Behavior
[What actually happens]

## Impact
- **Severity:** [Critical/High/Medium/Low]
- **Affected Users:** [Estimate]
- **Business Impact:** [Revenue/Trust/Conversion impact]

## Environment
- Platform: {brand}
- Category: {subtag}

## Possible Root Cause
[Initial hypothesis]

## Suggested Fix
[If obvious, suggest a fix approach]"""

    return generate_gpt_doc(prompt, f"You are a senior PM creating a {ticket_type} ticket.", max_tokens=2000)


def generate_multi_signal_prd(text_list, filename, brand="eBay"):
    combined = "\n\n".join(text_list)
    return generate_prd_docx(combined, brand, filename)


def _cluster_text_brand_and_meta(cluster_or_card):
    """Extract text, brand, and aggregated metadata from a cluster or card."""
    if isinstance(cluster_or_card, dict) and "quotes" in cluster_or_card:
        text = "\n\n".join(q.strip("- _") for q in cluster_or_card["quotes"])
        brand = cluster_or_card.get("brand", "eBay")
        meta_fields = {
            "Cluster Theme": cluster_or_card.get("theme", "Unknown"),
            "Problem Statement": cluster_or_card.get("problem_statement", ""),
            "Personas": ", ".join(cluster_or_card.get("personas", [])),
            "Sentiments": ", ".join(cluster_or_card.get("sentiments", [])),
            "Effort Levels": ", ".join(cluster_or_card.get("effort_levels", [])),
            "Topic Focus Tags": ", ".join(cluster_or_card.get("topic_focus_tags", [])),
            "Insight Count": str(cluster_or_card.get("insight_count", 0)),
            "Score Range": cluster_or_card.get("score_range", "N/A"),
        }
    else:
        items = cluster_or_card[:8] if isinstance(cluster_or_card, list) else []
        text = "\n\n".join(i.get("text", "") for i in items if i.get("text"))
        brand = items[0].get("target_brand", "eBay") if items else "eBay"
        personas = list({i.get("persona", "General") for i in items})
        sentiments = list({i.get("brand_sentiment", "Neutral") for i in items})
        topics = list({t for i in items for t in (i.get("topic_focus", []) or [])})
        meta_fields = {
            "Personas": ", ".join(personas),
            "Sentiments": ", ".join(sentiments),
            "Topic Focus Tags": ", ".join(topics[:6]),
            "Insight Count": str(len(items)),
        }
    return text, brand, meta_fields


def generate_cluster_prd_docx(cluster_or_card, filename):
    text, brand, meta_fields = _cluster_text_brand_and_meta(cluster_or_card)
    return generate_prd_docx(text, brand, filename, meta_fields=meta_fields)


def generate_cluster_brd_docx(cluster_or_card, filename):
    text, brand, meta_fields = _cluster_text_brand_and_meta(cluster_or_card)
    return generate_brd_docx(text, brand, filename, meta_fields=meta_fields)


def generate_cluster_prfaq_docx(cluster_or_card, filename):
    text, brand, meta_fields = _cluster_text_brand_and_meta(cluster_or_card)
    return generate_prfaq_docx(text, brand, filename, meta_fields=meta_fields)
